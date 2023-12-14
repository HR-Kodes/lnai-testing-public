from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views import View
import requests
import json
import urllib.parse
from django.http import HttpResponseBadRequest
import io
import urllib.request
import tempfile
from PIL import Image
from easyocr import Reader
from ultralytics.models.yolo import YOLO
from django.shortcuts import render
from django.http import HttpResponse
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from rest_framework.views import APIView
from rest_framework.response import Response
from django.views.generic import TemplateView
from docx import Document
from docx.shared import Pt
import openpyxl
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import requests
from rest_framework import status








class CompanyAvailabilityView(View):
    @csrf_exempt
    def dispatch(self, *args, **kwargs):
        return super().dispatch(*args, **kwargs)

    def get(self, request, *args, **kwargs):
        c_name = self.kwargs.get('c_name')  # Get the unique string name from the URL parameter
        if not c_name:
            return HttpResponseBadRequest("Missing 'c_name' parameter")

        print(f"Received company name: {c_name}")

        url = "https://www.mca.gov.in/bin/mca/login"
        headers = {
            "Origin": "www.mca.gov.in",
            "Content-Length": "231",
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/113.0.5672.127 Safari/537.36",
            "Content-Type": "application/x-www-form-urlencoded; charset=UTF-8",
            "Accept": "*/*",
            "X-Requested-With": "XMLHttpRequest",
            "Sec-Ch-Ua-Platform": "Windows"
        }
        data = "data=Ut8pBOc0RSM6iYqffqN1ovf7bobPRWJxrpoJRNCmK3GGtEoRKl3FZEf8xy36Iw3GHkoeOLH2Iw8tFFM6yB%2F6gZ4FgTyvicwoT%2FfDKJjUEuk9iiwbfqGT4bPey9LVPXwmJrnHl2AVXwPwjTY7BjtpUZ7CCRrYzb8G8Hs%2FTqLxp0oBRcfBtzGPnIVN9fMdkw%2Fj%2FSvMTL%2FNk05TbSN%2FPlQW%2FgdaTJRMuzmlIKp26ixHs8k%3D"
        response = requests.post(url, headers=headers, data=data)
        
        set_cookie_value = response.headers.get('Set-Cookie')
        start_index = set_cookie_value.find("session-token=") + len("session-token=")
        end_index = set_cookie_value.find(";", start_index)
        session_token = set_cookie_value[start_index:end_index]
        
        formData = "%7B%22formData%22%3A%7B%22companyType%22%3A%22New+Company+(Others)%22%2C%22companyClass%22%3A%22Private%22%2C%22companyCategory%22%3A%22Company+limited+by+shares%22%2C%22companySubCategory%22%3A%22Non-government+company%22%2C%22proposedName1%22%3A%22DUMBMET+PRIVATE+LIMITED%22%2C%22NICCode1%22%3A%2298200%22%2C%22Description1%22%3A%22Undifferentiated+service-producing+activities+of+private+households+for+own+use%22%2C%22NICCode2%22%3A%2298100%22%2C%22Description2%22%3A%22Undifferentiated+goods-producing+activities+of+private+households+for+own+use%22%2C%22formIntegrationId%22%3A%221%22%2C%22continueFlag%22%3A%22N%22%2C%22LLPIN%22%3A%22%22%2C%22NICCode3%22%3A%22%22%2C%22Description3%22%3A%22%22%2C%22proposedName2%22%3A%22%22%7D%2C%22formDescription%22%3A%22SPICE+PART+A%22%2C%22formName%22%3A%22Spice%2B+Part+A%22%2C%22formVersion%22%3A%221%22%2C%22userId%22%3A%22BIPULKUMARSINGH6690%40GMAIL.COM%22%2C%22integrationId%22%3A%221%22%2C%22prefill%22%3A%22false%22%2C%22status%22%3A%22Draft%2FPending+Submission%22%2C%22operation%22%3A%22Save%22%2C%22referenceNumber%22%3A%22%22%2C%22srn%22%3A%22%22%2C%22formId%22%3A%22%22%2C%22Approvedname%22%3A%22%22%2C%22serveAction%22%3A%22autocheck%22%7D"  # Your formData string here

        decoded_data = urllib.parse.unquote_plus(formData, encoding='utf-8')
        data_start = decoded_data.find('{')
        data_end = decoded_data.rfind('}') + 1
        data_str = decoded_data[data_start:data_end]
        data = eval(data_str)
        data['formData']['proposedName1'] = c_name
        encoded_data = urllib.parse.quote_plus(json.dumps(data), encoding='utf-8').replace('%20', '').replace('%28', '(').replace('%29', ')').replace('+%', '%')

        url = "https://www.mca.gov.in/bin/mca-gov/newSpiceA"
        headers = {
            "Host": "www.mca.gov.in",
            "Cookie": f"cookiesession1=678B28695C218253C321286001478935; alertPopup=true; session-token={session_token}; deviceId=1uytaas0tlbi",
            "Content-Type": "application/x-www-form-urlencoded; charset=UTF-8",
            "Origin": "https://www.mca.gov.in",
            "Referer": "https://www.mca.gov.in/content/mca/global/en/mca/e-filing/incorporation/spice.html"
        }
        payload = f"formData={encoded_data}"
        response = requests.post(url, headers=headers, data=payload)
        
        result = response.json()
        alert_list = result['validationResponse']['validationresposeBody']
        all_empty = all(not item['alertDescription'] for item in alert_list)

        print(f"Validation alert list: {alert_list}")
        
        if all_empty:
            print(f"Availability for {c_name}: available")
            availability_response = {"name": c_name, "availability": f"{c_name} is available"}
        else:
            print(f"Availability for {c_name}: not available")
            availability_response = {"name": c_name, "availability": f"{c_name } is not available"}

        return JsonResponse(availability_response)





class PincodeinfoView(View):
   
    @csrf_exempt
    def dispatch(self, *args, **kwargs):
        return super().dispatch(*args, **kwargs)

    def get(self, request, *args, **kwargs):
        pincode = self.kwargs.get('pincode')  # Get the pincode from the URL parameter
        if not pincode:
            return HttpResponseBadRequest("Missing 'pincode' parameter")

        print(f"Received pincode: {pincode}")



        url = "https://www.mca.gov.in/bin/mca/login"

        headers = {
            "Origin": "www.mca.gov.in",
            "Content-Length": "235",
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.5735.199 Safari/537.36",
            "Content-Type": "application/x-www-form-urlencoded; charset=UTF-8",
            "Accept": "*/*",
            "X-Requested-With": "XMLHttpRequest",
            "Sec-Ch-Ua-Platform": "Windows"
        }

        data = "data=Ut8pBOc0RSM6iYqffqN1ovf7bobPRWJxrpoJRNCmK3GGtEoRKl3FZEf8xy36Iw3GHkoeOLH2Iw8tFFM6yB%2F6gZ4FgTyvicwoT%2FfDKJjUEuk9iiwbfqGT4bPey9LVPXwmJrnHl2AVXwPwjTY7BjtpUZ7CCRrYzb8G8Hs%2FTqLxp0oBRcfBtzGPnIVN9fMdkw%2Fj%2FSvMTL%2FNk05TbSN%2FPlQW%2FgdaTJRMuzmlIKp26ixHs8k%3D"
        response = requests.post(url, headers=headers, data=data)
        set_cookie_value = response.headers.get('Set-Cookie')
        start_index = set_cookie_value.find("session-token=") + len("session-token=")
        end_index = set_cookie_value.find(";", start_index)
        session_token = set_cookie_value[start_index:end_index]
        url = 'https://www.mca.gov.in/content/forms/af/mca-aem-forms/form-fillip/fillip-main-form/form-fillip/jcr:content/guideContainer.af.dermis'

        headers = {
            'Host': 'www.mca.gov.in',
            'Cookie': f'cookiesession1=678B2869144E6B476733B93A4104896E; deviceId=4d23ogqkdbe; alertPopup=true; session-token{session_token}',
            'Content-Length': '544',
            'Sec-Ch-Ua': '',
            'Sec-Ch-Ua-Mobile': '?0',
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.5735.199 Safari/537.36',
            'Content-Type': 'application/x-www-form-urlencoded; charset=UTF-8',
            'Accept': 'text/plain, */*; q=0.01',
            'Csrf-Token': 'undefined',
            'X-Requested-With': 'XMLHttpRequest',
            'Sec-Ch-Ua-Platform': "",
            'Origin': 'https://www.mca.gov.in',
            'Sec-Fetch-Site': 'same-origin',
            'Sec-Fetch-Mode': 'cors',
            'Sec-Fetch-Dest': 'empty',
            'Referer': 'https://www.mca.gov.in/content/mca/global/en/mca/llp-e-filling/Fillip.html',
            'Accept-Encoding': 'gzip, deflate',
            'Accept-Language': 'en-US,en;q=0.9'
        }
        print("a")
        payload = {
            'functionToExecute': 'invokeFDMOperation',
            'formDataModelId': '/content/dam/formsanddocuments-fdm/mca-aem-forms/common-service/common-pin-info',
            "input": f'{{"Model0":{{"value":"{pincode}"}}}}',
            'operationName': 'POST /bin/mca-gov/commonpincode',
            'guideNodePath': '/content/forms/af/mca-aem-forms/form-fillip/fillip-main-form/block2/jcr:content/guideContainer/rootPanel/items/panel_1379931518_cop/items/panel/items/panel/items/panel_702814714/items/panel/items/guidetextbox_copy_11_1386963699'
        }

        response = requests.post(url, headers=headers, data=payload)

        response_dict = json.loads(response.text)
        return JsonResponse(response_dict)
    



class DetectObjectsAndTextView(View):


    @csrf_exempt
    def dispatch(self, *args, **kwargs):
        return super().dispatch(*args, **kwargs)

    def post(self, request):
        try:
            data = json.loads(request.body)
            image_url = data.get('imageurl', '')

            if not image_url:
                return JsonResponse({'error': 'Image URL not provided'}, status=400)

            reader = Reader(['en'])
            model = YOLO('D:/office/spice backend/django backend/CompanyAvailabilityProject/companyavailabilityapp/best.pt')  # Update with the actual path to your YOLO model checkpoint

            t = model.predict(image_url, save=True, imgsz=640, conf=0.1, save_txt=True, save_conf=True, show_labels=False)

            img_paths = []
            label_paths = []
            img_name = image_url.split('/')[-1]

            for pred in t:
                labels = pred.names

                img_path = pred.save_dir + '/' + img_name
                label_path = pred.save_dir + '/labels/' + '.'.join(img_name.split('.')[:-1]) + '.txt'

                img_paths.append(img_path)
                label_paths.append(label_path)

            path = io.BytesIO(urllib.request.urlopen(image_url).read())
            detect_list = []

            for x in range(len(img_paths)):
                with open(label_paths[x], 'r') as file:
                    label = file.readlines()
                img = Image.open(path)

                height = img.height
                width = img.width
                detect_dict = {}

                for i in label:
                    value = [float(a) for a in i.split()]
                    sw = int(width * float(value[1] - value[3] / 2))
                    ew = int(width * float(value[1] + value[3] / 2))
                    sh = int(height * float(value[2] - value[4] / 2))
                    eh = int(height * float(value[2] + value[4] / 2))

                    if labels[int(value[0])] not in detect_dict:
                        detect_dict[labels[int(value[0])]] = [[img.crop((sw, sh, ew, eh)), value[5]]]
                    else:
                        detect_dict[labels[int(value[0])]].append([img.crop((sw, sh, ew, eh)), value[5]])

                detect_list.append(detect_dict)

            detect_list1 = [x.copy() for x in detect_list]

            for f, i in enumerate(detect_list):
                for x, y in i.items():
                    tl = []
                    highest_confidence = 0.0
                    best_result = ""

                    for t in y:
                        temp_image_path = tempfile.NamedTemporaryFile(suffix='.jpg').name
                        t[0].save(temp_image_path, format='JPEG')
                        results = reader.readtext(temp_image_path)

                        for detection in results:
                            text = detection[1]
                            confidence = t[1]

                            if confidence > highest_confidence:
                                highest_confidence = confidence
                                best_result = text

                        tl = [best_result, confidence]

                    detect_list1[f][x] = tl

            return JsonResponse(detect_list1, safe=False)

        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON data'}, status=400)
        



class AttendanceSheet(APIView):
    def post(self, request, *args, **kwargs):
        type_of_meeting = request.data.get('type_of_meeting')
        serial_number = request.data.get('serial_number')
        date_of_meeting = request.data.get('date_of_meeting')
        time_of_meeting = request.data.get('time_of_meeting')
        mode_of_meeting = request.data.get('mode_of_meeting')
        company_name = request.data.get('company_name')
        chairman_name = request.data.get('chairman_name')
        chairman_din = request.data.get('chairman_din')
        place_of_meeting = request.data.get('place_of_meeting')
        directors_count = int(request.data.get('directors_count', 0))

        doc = Document()
        heading = doc.add_heading("Attendance Sheet", level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        print("Date of Meeting:", date_of_meeting)
        print("Chairman Name:", chairman_name)
        print("Directors Count:", directors_count)
        print("Place of Meeting:", place_of_meeting)

        # Create the table for director details
        directors_table = doc.add_table(rows=1, cols=4)
        directors_table.style = "Table Grid"

        # Set the column widths
        column_widths = [1, 3, 3, 3]
        for i, width in enumerate(column_widths):
            directors_table.cell(0, i).width = width

        # Add the column headers for director details
        directors_headers = directors_table.rows[0].cells
        directors_headers[0].text = "Sr. No."
        directors_headers[1].text = "Name of Directors"
        directors_headers[2].text = "Designation"
        directors_headers[3].text = "Mode of Presence"

        for i in range(directors_count):
            sr_no = str(i + 1)
            director_data = request.data.get(f'director_{sr_no}', {})

            name = director_data.get('name', '')
            designation = director_data.get('designation', '')
            din_number = director_data.get('din_number', '')
            mode_of_presence = director_data.get('mode_of_presence', '')

            row = directors_table.add_row().cells
            row[0].text = sr_no
            row[1].text = name

            # Add designation and DIN number in the same cell
            cell_3_paragraph = row[2].paragraphs[0]
            cell_3_paragraph.text = designation
            cell_3_paragraph.add_run("\n")  # Add a newline between designation and DIN number
            cell_3_paragraph.add_run(din_number)

            row[3].text = mode_of_presence

        # Add a blank paragraph for spacing
        doc.add_paragraph()

        # Add the company's name and chairman's details
        paragraph = doc.add_paragraph()
        paragraph.add_run("For and on behalf of the Board of Directors,").bold = True
        paragraph.add_run("\n" + company_name)

        # Add a blank paragraph for spacing
        doc.add_paragraph()

        # Add the line for signature
        doc.add_paragraph("_____________________________")

        # Add the chairman's details
        doc.add_paragraph(chairman_name)
        doc.add_paragraph("Chairman of the meeting")
        doc.add_paragraph(f"DIN: {chairman_din}")

        # Add a blank paragraph for spacing
        doc.add_paragraph()

        # Add the date and place of the meeting
        doc.add_paragraph(f"Date: {date_of_meeting}")
        doc.add_paragraph(f"Place: {place_of_meeting}")

        # Save the document
        doc.save("attendance_sheet.docx")

        # Create a new document
        doc = Document()
        heading = doc.add_heading("Attendance Sheet", level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=attendance_sheet.docx'
        file_path = f"directors_table_{directors_count}_directors.docx"
        doc.save(file_path)
        return Response({'success': True, 'file_path': file_path})


class DirectorList(APIView):
    def post(self, request, *args, **kwargs):

        num_directors = int(request.data.get('num_directors', 0))

        if num_directors < 1:
            return Response({'error': 'Invalid input. Please enter a positive integer.'}, status=400)

        doc = Document()
        heading_paragraph = doc.add_paragraph()
        heading_run = heading_paragraph.add_run("LIST OF DIRECTORS")
        heading_run.bold = True
        # heading_run.font.size = Pt(14)
        heading_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        headers = [request.data.get('din_header', 'DIN/PAN/DPIN'),
                   request.data.get('name_header', 'Full Name'),
                   request.data.get('designation_header', 'Designation'),
                   request.data.get('address_header', 'Address')]
        
        
       
        header_cells = table.rows[0].cells
        for i in range(len(headers)):
            header_cells[i].text = headers[i]

        for i in range(num_directors):
            row = table.add_row().cells
            row[0].text = request.data.get(f'din_{i}', '')
            row[1].text = request.data.get(f'full_name_{i}', '')
            row[2].text = request.data.get(f'designation_{i}', '')
            row[3].text = request.data.get(f'address_{i}', '')

        file_path = f"directors_table_{num_directors}_directors.docx"
        doc.save(file_path)

        return Response({'success': True, 'file_path': file_path})




class Minutemeet(APIView):
    def post(self, request, *args, **kwargs):
        data = request.data

        # Extract data from the request
        board_meeting_number = data.get('board_meeting_number')
        company_name = data.get('company_name')
        date_of_meeting = data.get('date_of_meeting')
        location_of_meeting = data.get('location_of_meeting')
        directors_count = int(data.get('directors_count'))

        # Create a new document
        doc = Document()

        # Your existing code here...
        # You can reuse the logic from your original code
        # Make sure to replace user input with data from the request

        # Save the document to a BytesIO object
        from io import BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=board_meeting_minutes.docx'
        response.write(output.read())

        return response








class RelatedDirectorCompanies(APIView):
    
    def get(self, request, cin, format=None):
        try:
            if not cin:
                return Response({"error": "CIN is required"}, status=status.HTTP_400_BAD_REQUEST)
            
            url = f"https://api.finanvo.in/company/directors?CIN={cin}"
            headers = {
                "Origin": "https://web.compdata.in",
                
            }

            response = requests.get(url, headers=headers)

            if response.status_code == 200:
                return JsonResponse(response.json(), status=status.HTTP_200_OK)
            else:
                return JsonResponse({"error": f"Request failed with status code: {response.status_code}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class Cindata(APIView):
    
    def get(self, request, cin, format=None):
        try:
            if not cin:
                return Response({"error": "CIN is required"}, status=status.HTTP_400_BAD_REQUEST)
            
            url = f"https://api.finanvo.in/company/profile?CIN={cin}"
            headers = {
                "Host": "api.finanvo.in",
                "App-Origin": "https://web.compdata.in",
                "Origin": "https://web.compdata.in",
                "Referer": "https://web.compdata.in/",
            }

            response = requests.get(url, headers=headers)

            if response.status_code == 200:
                return JsonResponse(response.json(), status=status.HTTP_200_OK)
            else:
                return JsonResponse({"error": f"Request failed with status code: {response.status_code}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)















# class GenerateNoticeView(APIView):
#     def replace_variables(self, proposed_solution, data):
#         with open(data, 'r') as file:
#             json_data = json.load(file)

#         for key, value in json_data.items():
#             variable = f"{{{key}}}"
#             if variable in proposed_solution:
#                 proposed_solution = proposed_solution.replace(variable, str(value))
#         return proposed_solution


#     def post(self, request, *args, **kwargs):
            
#             type_of_meeting = request.data.get('type_of_meeting')
#             serial_number = request.data.get('serial_number')
#             date_of_meeting = request.data.get('date_of_meeting')
#             time_of_meeting = request.data.get('time_of_meeting')
#             company_name = request.data.get('company_name')
#             exact_address = request.data.get('exact_address')
#             place_of_meet = request.data.get('place_of_meet')
#             chairman_name = request.data.get('chairman_name')
#             DIN_number= request.data.get('DIN_number')
#             director_name = request.data.get('director_name')
#             agenda_item= request.data.get('agenda_item')
            

            
#             doc = Document()





#             # Create the letterhead section
#             #letterhead = doc.sections[0].header
#             #letterhead.is_linked_to_previous = False

#             # Add the heading to the letterhead
#             #heading = letterhead.paragraphs[0].add_run("Letterhead")
#             #heading.bold = False
#             #heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

#             # Print the board meeting number
#             doc.add_paragraph(f"Board Meeting Number: {serial_number}")

#             # Add the notice of the meeting
#             notice_paragraph = doc.add_paragraph()
#             notice_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#             # Add the bold "Notice of Meeting" text
#             notice_run = notice_paragraph.add_run("Notice of Board Meeting")
#             notice_run.bold = True
#             notice_run.font.size = Pt(14)

#             doc.add_paragraph("To,\n"
#                             "The Board of Directors,\n"
#                             f"{company_name},\n"
#                             f"{exact_address}")

#             # Add the salutation
#             doc.add_paragraph("Respected Sir/Madam,")

#             # Add the notice of meeting
#             paragraph = doc.add_paragraph("Notice is hereby given that the meeting of the Board of Directors of the Company "
#                                         + company_name + " "
#                                         + "(the 'Board') scheduled to be held on " + date_of_meeting + " at " + time_of_meeting + " "
#                                         + "at the registered office of the company situated " + exact_address + ", India "
#                                         + "to transact the below-mentioned business:")

#             # Initialize the agenda counter
#             agenda_counter = 1

#             # Create a list to store the agenda items and proposals
#             agenda_list = []

#             # Iterate through multiple agenda items
#             while True:
#                 # Prompt the user for the agenda item
                
                
#                 if agenda_item.lower() == "done":
#                     break

#                 # Load the Excel file and retrieve the corresponding proposal
#                 workbook = openpyxl.load_workbook("Untitled 2 2 2.xlsx")
#                 worksheet = workbook["Sheet1"]

#                 # Search for the corresponding proposal based on the agenda item
#                 proposal = None
#                 for row in worksheet.iter_rows(values_only=True):
#                     if row[0] == agenda_item:
#                         proposal = row[1]
#                         break

#                 # Replace variables in the proposal
#                 proposal = replace_variables(proposal, "data.json")
                
#                 # Print the agenda item with the counter number
#                 agenda_text = f"{agenda_counter}. {agenda_item}"
                
#                 # Append the agenda item and proposal to the list
#                 agenda_list.append((agenda_text,proposal))
                
#                 # Increment the agenda counter
#                 agenda_counter += 1

#             # Create a new paragraph in the document for the agenda list
#             doc.add_paragraph(" ")
#             agenda_list_paragraph = doc.add_paragraph()

#             # Add each agenda item to the document
#             for agenda_item, _ in agenda_list:
#                 agenda_list_paragraph.add_run(agenda_item + "\n")
#                 #agenda_list_paragraph.add_run(proposal + "\n\n")
            
#             #Add the additional paragraph after the agenda list
#             doc.add_paragraph("Kindly make it convenient to attend the meeting. Please submit leave of absence in case you are not in a position to attend the meeting.")

#             # Add the additional paragraph after the agenda list
#             additional_paragraph = doc.add_paragraph("By and on behalf of \n")
#             additional_paragraph.add_run(company_name).bold = True
#             additional_paragraph.add_run("\n\n_____________")
#             additional_paragraph.add_run("\n\n")
#             additional_paragraph.add_run(chairman_name)
#             additional_paragraph.add_run("\nDirector")
#             additional_paragraph.add_run("\nDIN: " + DIN_number)
#             additional_paragraph.add_run("\n\nDate: " + date_of_meeting)
#             additional_paragraph.add_run("\nPlace: " + exact_address)

#             # Copy of the notice to:
#             doc.add_paragraph("\nCopy of the notice to:\n")
#             dir_1_present =  input("Name of the first director present ")
#             dir_2_present =  input("Name of the second director present ")
#             doc.add_paragraph(f"{dir_1_present}")
#             doc.add_paragraph(f"{dir_2_present}")

#             #enter the chairman of the meet
#             chairman_present = input("enter the chairman present ")
#             din_of_chairman = input("enter the din of the chairman ")
#             # Add the NOTES TO AGENDA section
#             notes_to_agenda = f"NOTES TO AGENDA: {type_of_meeting} DATED {date_of_meeting}:"
#             notes_to_agenda_paragraph = doc.add_paragraph("\n" + notes_to_agenda)
#             notes_to_agenda_paragraph.bold = True

#             # Iterate through each agenda item and proposal
#             for agenda_text, proposal in agenda_list:
#                 item_paragraph = doc.add_paragraph("ITEM NO. " + str(agenda_counter) + ": " + agenda_text)
#                 item_paragraph.add_run("\n\n" + proposal)
                
#                 # Increment the agenda counter
#                 agenda_counter += 1

                
#             # Add the "By and on behalf of" section
#             by_on_behalf_paragraph = doc.add_paragraph("By and on behalf of\n")
#             by_on_behalf_paragraph.add_run(company_name).bold = True
#             by_on_behalf_paragraph.add_run("\n\n_____________")
#             by_on_behalf_paragraph.add_run("\n\n")
#             by_on_behalf_paragraph.add_run("Chairman of the meeting")
#             by_on_behalf_paragraph.add_run("\n" + director_name)
#             by_on_behalf_paragraph.add_run("\nDIN: " + DIN_number)
#             by_on_behalf_paragraph.add_run("\n\nDate: " + date_of_meeting)
#             by_on_behalf_paragraph.add_run("\nPlace: " +place_of_meet)

#             # Save the document     
            

#             # Create a new document

           

#             # Save the document
#             doc.save("notice_with_agenda_list.docx")

#             # Provide the file for download
#             with open("notice_with_agenda_list.docx", "rb") as docx_file:
#                 response = HttpResponse(docx_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
#                 response['Content-Disposition'] = 'attachment; filename=notice_with_agenda_list.docx'
#                 return response

#             return render(request, "generate_notice.html")



