from azure.core.credentials import AzureKeyCredential
from azure.ai.formrecognizer import DocumentAnalysisClient

import re

from docx.api import Document

import os
import subprocess
from datetime import datetime



endpoint = os.environ.get('MY_ENDPOINT')
key = os.environ.get('MY_SECRET_KEY')
    

def analyze_read(document_path):
    document_analysis_client = DocumentAnalysisClient(
        endpoint=endpoint, credential=AzureKeyCredential(key)
    )

    with open(document_path, "rb") as f:
        poller = document_analysis_client.begin_analyze_document(
            "prebuilt-read", document=f
        )
        result = poller.result()
    return result.content


def extract_address(text):
    match = re.search(r'\d{6}', text)
    if match:
        pincode_index = match.start()
        colon_index = text.rfind(':', 0, pincode_index)
        if colon_index != -1:
            result = text[colon_index + 1:pincode_index + 6].strip()
            return result
        else:
            print("No matching pattern found")
    else:
        print("No matching pattern found")


def extract_aadhar_details(aadhar_text):
    non_english_regex = r'[^\x00-\x7F]+'
    filtered_text = re.sub(non_english_regex, '', aadhar_text)
    filtered_text = ' '.join(filtered_text.split()).replace('\n,', ',').replace(',\n', ',').replace('\n', ' ')
    pattern = r"\d{4}\s+\d{4}\s+\d{4}"
    aadhar_id = re.findall(pattern, filtered_text)[0]
    if len(aadhar_text) < 700:
        match = re.search(r"INDIA (.*?) / DOB", filtered_text)
        if match:
            name = match.group(1)
        else:
            print("couldn't find name")
            name = ''

        match = re.search(r"Address: (.*?\d{6})", filtered_text)
        if match:
            personal_address = match.group(1)
            personal_address = re.split('[,.]', personal_address)
            if 'S/O' in personal_address[0]:
                personal_address[0] = personal_address[0].replace('S/O:', '')
                personal_address.pop(0)
            personal_address = "".join(i+',\n' for i in personal_address).strip(',\n')
        else:
            print("couldn't find the address")
            personal_address = ''
    else:
        match = re.search(r'INDIA>(.*?)/ Year', filtered_text)
        if match:
            name = match.group(1)
        else:
            print("couldn't find name")
            name = ''

        match = re.search(r'Address:(.*?(\d{6}))', filtered_text)
        if match:
            personal_address = match.group(1).replace(", ,", ",")
            personal_address = re.split('[,.]', personal_address)
            personal_address = "".join(i+',\n' for i in personal_address).strip(',\n')
        else:
            print("No address found.")
            personal_address = ''
    return {'personal_name': name, 'aadhar_id': aadhar_id, 'personal_address': personal_address}


def combine_word_documents(doc1, doc2):
    merged_document = Document()

    sub_doc = Document(doc2)
    doc1.add_page_break()
    for element in doc1.element.body:
            merged_document.element.body.append(element)
    
    for element in sub_doc.element.body:
            merged_document.element.body.append(element)
    
    return merged_document


def create_form9_pdfs(number_of_partners, form_9_document, data, llp_address):
    
    today_date = datetime.today().strftime('%d-%m-%Y')
    replace_tg_in_paragraph(form_9_document, '<date>', today_date)
    replace_tg_in_paragraph(form_9_document, '<place>', 'Bangalore')
    replace_tg_in_paragraph(form_9_document, '<llp_address>', llp_address)
    replace_tg_in_paragraph(form_9_document, '<Designated_Partner_name>', data["personal_name"])

    replace_tags_from_table(form_9_document, '<Designated_Partner_name>', data["personal_name"], 1, 1)
    replace_tags_from_table(form_9_document, '<Designated_Partner_address>', data["personal_address"].lower(), 3, 1)

    if number_of_partners > 0:
        return combine_word_documents(form_9_document, "Form_9_Jay.docx")
    else:
        form_9_document.save("updated_form9.docx")
        subprocess.call(['unoconv', '-f', 'pdf', 'updated_form9.docx'])


def create_subscriber_sheet_pdf(partners, Subscriber_Sheet, data, witness_data):

    today_date = datetime.today().strftime('%d-%m-%Y')
    replace_tg_in_paragraph(Subscriber_Sheet, "<date>", today_date)
    replace_tg_in_paragraph(Subscriber_Sheet, "<place>", "Bangalore")
    replace_tags_from_table(Subscriber_Sheet, "<name_of_witness>", witness_data["personal_name"], number_of_partners-partners, 3)

    if partners > 0:
        table = Subscriber_Sheet.tables[0]
        new_row = table.add_row()
        previous_row = table.rows[-2]
        new_row.cells[0].text = previous_row.cells[0].text
        new_row.cells[1].text = previous_row.cells[1].text
        replace_tags_from_table(Subscriber_Sheet, "<Designated_Partner_name>", data["personal_name"]+'\n\n\n', number_of_partners-partners, 0)
        return Subscriber_Sheet
    else:
        replace_tags_from_table(Subscriber_Sheet, "<Designated_Partner_name>", data["personal_name"]+'\n\n\n', number_of_partners-partners, 0)
        Subscriber_Sheet.save("updated_Subscriber_Sheet.docx")
        subprocess.call(['unoconv', '-f', 'pdf', 'updated_Subscriber_Sheet.docx'])


def replace_tg_in_paragraph(document, tag, replace_text):
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if tag in run.text:
                replaced_run = run._element
                replaced_run.text = run.text.replace(tag, replace_text)
                for prop in run._element.rPr:
                    if prop.tag.endswith('b'):
                        replaced_run.rPr.append(prop)
                paragraph._p.replace(run._element, replaced_run)

    return document

def replace_tags_from_table(document, tag, replace_text, row, column):
    cell = document.tables[-1].cell(row, column) 
    cell_text = cell.text
    try:
        font_size = cell.paragraphs[0].runs[0].font.size
        new_text = cell_text.replace(tag, replace_text)
        cell.text = ""
        new_run = cell.paragraphs[0].add_run(new_text)
        new_run.font.size = font_size
    except:
        pass
    return document


def create_part_b_pdf( number_of_partners, document, data):

    replace_tg_in_paragraph(document, '<Designated_Partner_name>',data["personal_name"])
    address = data["personal_address"]
    address = address.replace("\n", " " )
    data['personal_address'] = address
    replace_tg_in_paragraph(document, '<Designated_Partner_address>',data["personal_address"])

    if number_of_partners > 0:
        return combine_word_documents(document, 'Part_B_Statement.docx')
        
    else:
        part_b.save("updated_Part_B_Statement.docx")
        subprocess.call(['unoconv', '-f', 'pdf', 'updated_Part_B_Statement.docx'])


if __name__ == "__main__":
    number_of_partners = 2

    aadhar_files = ['Aadhaar_letter_large-3801909141.jpg', 'short_form_aadhar.jpg']
    witness_aadhar = 'short_form_aadhar.jpg'

    ciie_form = 'ventu_ciie.pdf'
    llp_address = extract_address(analyze_read(ciie_form))
    llp_address = re.split('[,.]', llp_address)
    llp_address = "".join(i+',\n' for i in llp_address).strip(',\n')

    witness_data = extract_aadhar_details(analyze_read(witness_aadhar))

    partners_data = []
    for Files in aadhar_files:
        partners_data.append(extract_aadhar_details(analyze_read(Files)))


    form_9 = Document('Form_9_Jay.docx')
    part_b = Document('Part_B_Statement.docx')
    Subscriber_Sheet_doc = Document('Subscriber_Sheet.docx')

    for i in reversed(range(number_of_partners)):
        form_9 = create_form9_pdfs(i, form_9, partners_data[-1-i], llp_address)
        part_b = create_part_b_pdf(i, part_b, partners_data[-1-i])
        Subscriber_Sheet_doc = create_subscriber_sheet_pdf(i, Subscriber_Sheet_doc, partners_data[-1-i], witness_data)
